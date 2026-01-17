
from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from .types import Alarm, ObservationSegment, PeopleCountSegment, Report, Session
from .utils_time import format_ts


REPORT_VERSION = "v1.2"

STRINGS = {
    "en": {
        "report_title": "Sampling Report",
        "section_cover_summary": "Cover + Summary",
        "summary_intro": (
            "This report summarizes the offline video analysis results. It includes sampling sessions, "
            "people count changes, and detected alarms."
        ),
        "key_findings_title": "Key Findings / 关键发现 (Top3)",
        "scope_title": "Scope",
        "scope_method": "Method: Offline analysis",
        "scope_outputs": (
            "Outputs: Sampling sessions, people change log, people count segments, observation segments, timeline"
        ),
        "scope_alarm_note": "Alarm list is shown only when alarms exist.",
        "section_video_algo": "Video & Algorithm Info",
        "video_algo_intro": "All models in the pipeline participate in inference and decision outputs.",
        "video_algo_params_intro": "Key parameters are listed in the table below.",
        "section_sampling_overview": "Sampling Overview",
        "sessions_intro": (
            "A sampling session (sampling operation segment) is defined when the sampling hole is OPEN and a "
            "sampling action is detected; sessions are categorized as \"Blocked\" or \"Unblocked\". Crew count is "
            "monitored and reported as change events and stabilized segments."
        ),
        "section_alarm_list": "Alarm List",
        "alarms_intro": "This section lists alarms detected during sampling sessions.",
        "section_people_changes": "Crew Count Change Log",
        "people_changes_intro_en": (
            "Crew count changes are recorded only after the new count stays stable for at least 2 seconds; "
            "evidence range is +/-5 seconds around the change time."
        ),
        "section_people_segments": "Crew Count Segments",
        "people_segments_intro_en": "Segments represent continuous ranges of the stabilized crew count.",
        "section_observation": "Observation Segments / 观察区间（OPEN但无采样）",
        "observation_intro_en": (
            "Observation segments indicate OPEN state without detected sampling; only people count changes are "
            "recorded and no crew compliance judgment is made."
        ),
        "section_timeline": "Timeline",
        "timeline_intro": "Timeline lists sampling sessions, people change events, and alarms for quick lookup.",
        "section_disclaimer": "Disclaimer",
        "disclaimer": (
            "This report is generated from offline video analysis. Results assist review and should be verified."
        ),
        "glossary_title": "Appendix: Glossary / 附录：术语表",
        "alarm_explain_crew_under": (
            "Crew count dropped below 2 during sampling session {session_id} and lasted {duration_s:.1f}s "
            "(exceeding the grace window {T_people_grace_s:.1f}s). This is considered \"Crew Understaffed\"."
        ),
        "alarm_explain_crew_over": (
            "Crew count exceeded 2 during sampling session {session_id} and lasted {duration_s:.1f}s. "
            "This is considered \"Crew Overstaffed\"."
        ),
        "alarm_explain_unblocked_insertion": (
            "During sampling session {session_id}, an unblocked insertion condition (OPEN + sampling + "
            "no_blocking) accumulated to {actual_s:.1f}s and triggered the threshold {threshold_s:.1f}s at "
            "{trigger_ts}. This indicates sampling without blocking."
        ),
        "alarm_explain_sampling_too_short": (
            "Sampling session {session_id} duration is {duration_s:.1f}s, which is below the minimum required "
            "{T_sampling_min_s:.0f}s. This triggers \"Sampling Too Short\"."
        ),
    },
    "zh": {
        "report_title": "采样报告",
        "section_cover_summary": "封面 + 总结",
        "summary_intro": "本报告为离线视频智能分析结果，汇总采样会话、人数变化记录与报警信息。",
        "key_findings_title": "关键发现 Top3",
        "scope_title": "分析范围",
        "scope_method": "方式：离线分析",
        "scope_outputs": "输出包含：采样会话、人数变化记录、人数稳定区间、观察区间、时间轴",
        "scope_alarm_note": "报警清单仅在存在报警时展示。",
        "section_video_algo": "视频与算法信息",
        "video_algo_intro": "算法链路中所有模型均参与推理与决策输出。",
        "video_algo_params_intro": "关键规则参数如下表所示。",
        "section_sampling_overview": "采样总览",
        "sessions_intro": (
            "采样会话（采样作业段）的定义为：采样孔处于开启（OPEN）且检测到采样行为；会话按封堵状态分为“封堵采样”"
            "与“无封堵采样”。采样会话期间监控人数变化，并以变化记录与稳定区间形式呈现。"
        ),
        "section_alarm_list": "报警清单",
        "alarms_intro": "本节列出采样会话期间检测到的报警信息。",
        "section_people_changes": "人数变化记录",
        "people_changes_intro_zh": "人数变化按 2 秒稳定去抖记录；复核范围为变化时刻前后各 5 秒。",
        "section_people_segments": "人数阶段区间",
        "people_segments_intro_zh": "人数阶段为去抖后的稳定区间，用于查看整体人数分布。",
        "section_observation": "Observation Segments / 观察区间（OPEN但无采样）",
        "observation_intro_zh": "观察区间表示开启但未检测到采样，仅记录人数变化，不做人数合规判定。",
        "section_timeline": "时间轴页",
        "timeline_intro": "时间轴页列出采样会话、人数变化与报警事件，便于快速定位。",
        "section_disclaimer": "免责声明",
        "disclaimer": "本报告为离线视频智能分析结果，结果用于辅助复核，需结合视频人工确认。",
        "glossary_title": "Appendix: Glossary / 附录：术语表",
        "alarm_explain_crew_under": (
            "在采样会话 {session_id} 期间，人数少于2人持续 {duration_s:.1f} 秒（超过容忍窗口 "
            "{T_people_grace_s:.1f} 秒），判定为“擅自离岗（人数不足）”。"
        ),
        "alarm_explain_crew_over": "在采样会话 {session_id} 期间，人数超过2人持续 {duration_s:.1f} 秒，判定为“串岗（人数超员）”。",
        "alarm_explain_unblocked_insertion": (
            "在采样会话 {session_id} 期间，检测到“无封堵伸入”（开启 + 采样 + 无封堵）条件累计达到 {actual_s:.1f} 秒，"
            "并在 {trigger_ts} 达到阈值 {threshold_s:.1f} 秒触发报警，表示存在无封堵采样风险。"
        ),
        "alarm_explain_sampling_too_short": (
            "采样会话 {session_id} 持续时长为 {duration_s:.1f} 秒，小于最短时长要求 {T_sampling_min_s:.0f} 秒，触发"
            "“采样时长不足”报警。"
        ),
    },
}


def _format_duration_s(value: float) -> str:
    return str(int(round(value)))


def _session_rows(sessions: List[Session]) -> List[List[str]]:
    rows = []
    for session in sessions:
        rows.append(
            [
                str(session.session_id),
                session.session_type,
                format_ts(session.start_ts_s),
                format_ts(session.end_ts_s),
                _format_duration_s(session.duration_s),
            ]
        )
    return rows


def _observation_rows(segments: List[ObservationSegment]) -> List[List[str]]:
    rows = []
    for idx, seg in enumerate(segments, start=1):
        rows.append(
            [
                str(idx),
                format_ts(seg.start_ts_s),
                format_ts(seg.end_ts_s),
                _format_duration_s(seg.duration_s),
            ]
        )
    return rows


def _context_label(flag: Optional[bool]) -> str:
    if flag is None:
        return "-"
    return "In sampling session / 采样期间" if flag else "Out of sampling session / 非采样期间"


def _people_change_rows(changes) -> List[List[str]]:
    rows = []
    for change in changes:
        evidence_start = max(0.0, change.change_ts_s - 5.0)
        evidence_end = change.change_ts_s + 5.0
        rows.append(
            [
                format_ts(change.change_ts_s),
                f"{change.from_count}→{change.to_count}",
                _context_label(change.context_in_session),
                f"{format_ts(evidence_start)} - {format_ts(evidence_end)}",
            ]
        )
    return rows


def _normalize_people_segments(segments: List[PeopleCountSegment]) -> List[PeopleCountSegment]:
    normalized: List[PeopleCountSegment] = []
    for seg in segments:
        rounded = int(round(seg.duration_s))
        if rounded <= 0:
            continue
        if normalized:
            last = normalized[-1]
            if (
                last.people_count == seg.people_count
                and last.context_in_session == seg.context_in_session
                and abs(last.end_ts_s - seg.start_ts_s) < 1e-6
            ):
                last.end_ts_s = seg.end_ts_s
                last.duration_s = max(0.0, last.end_ts_s - last.start_ts_s)
                continue
        normalized.append(
            PeopleCountSegment(
                start_ts_s=seg.start_ts_s,
                end_ts_s=seg.end_ts_s,
                duration_s=seg.duration_s,
                people_count=seg.people_count,
                context_in_session=seg.context_in_session,
            )
        )
    return normalized


def _people_segment_rows(segments: List[PeopleCountSegment]) -> List[List[str]]:
    rows = []
    for seg in _normalize_people_segments(segments):
        rows.append(
            [
                format_ts(seg.start_ts_s),
                format_ts(seg.end_ts_s),
                _format_duration_s(seg.duration_s),
                str(seg.people_count),
                _context_label(seg.context_in_session),
            ]
        )
    return rows


def _explain_alarm(
    alarm: Alarm,
    *,
    lang: str,
    report: Report,
    session_id: str,
    duration_s: float,
    evidence_start_s: float,
    evidence_end_s: float,
) -> str:
    strings = STRINGS[lang]
    if alarm.alarm_type in ("CREW_UNDER", "CREW_UNDER_2"):
        return strings["alarm_explain_crew_under"].format(
            session_id=session_id,
            duration_s=duration_s,
            T_people_grace_s=report.config.people_grace_s,
        )
    if alarm.alarm_type in ("CREW_OVER", "CREW_OVER_2"):
        return strings["alarm_explain_crew_over"].format(
            session_id=session_id,
            duration_s=duration_s,
        )
    if alarm.alarm_type == "UNBLOCKED_INSERTION":
        trigger_ts = "-" if alarm.trigger_ts_s is None else format_ts(alarm.trigger_ts_s)
        return strings["alarm_explain_unblocked_insertion"].format(
            session_id=session_id,
            actual_s=duration_s,
            threshold_s=report.config.unblocked_alarm_s,
            trigger_ts=trigger_ts,
        )
    if alarm.alarm_type == "SAMPLING_TOO_SHORT":
        return strings["alarm_explain_sampling_too_short"].format(
            session_id=session_id,
            duration_s=duration_s,
            T_sampling_min_s=report.config.sampling_min_s,
        )
    return f"{alarm.alarm_type} ({format_ts(evidence_start_s)}-{format_ts(evidence_end_s)})"


def _infer_session_id(alarm: Alarm, sessions: List[Session]) -> str:
    if alarm.session_id is not None:
        return str(alarm.session_id)
    for session in sessions:
        if alarm.start_ts_s >= session.start_ts_s and alarm.start_ts_s <= session.end_ts_s:
            return str(session.session_id)
    return "-"


def _build_key_findings(report: Report) -> List[tuple[str, str]]:
    findings: List[tuple[str, str]] = []
    segments = _normalize_people_segments(report.people_count_segments)
    if segments:
        max_seg = max(segments, key=lambda s: (s.people_count, s.duration_s))
        findings.append(
            (
                f"Max people count {max_seg.people_count} during {format_ts(max_seg.start_ts_s)}-{format_ts(max_seg.end_ts_s)}.",
                f"最大人数 {max_seg.people_count}，时间段 {format_ts(max_seg.start_ts_s)}-{format_ts(max_seg.end_ts_s)}。",
            )
        )
    zero_seg = next((seg for seg in segments if seg.people_count == 0), None)
    if zero_seg:
        findings.append(
            (
                f"People count 0 during {format_ts(zero_seg.start_ts_s)}-{format_ts(zero_seg.end_ts_s)}.",
                f"人数为0时间段 {format_ts(zero_seg.start_ts_s)}-{format_ts(zero_seg.end_ts_s)}。",
            )
        )
    else:
        findings.append(("No people_count=0 segment detected.", "未检测到人数为0的区间。"))

    if report.people_count_change_events:
        change = max(
            report.people_count_change_events,
            key=lambda c: abs(c.to_count - c.from_count),
        )
        evidence_start = max(0.0, change.change_ts_s - 5.0)
        evidence_end = change.change_ts_s + 5.0
        findings.append(
            (
                f"Largest change {change.from_count}→{change.to_count} at {format_ts(change.change_ts_s)} "
                f"(evidence {format_ts(evidence_start)}-{format_ts(evidence_end)}).",
                f"最大变化 {change.from_count}→{change.to_count} 发生于 {format_ts(change.change_ts_s)} "
                f"（证据 {format_ts(evidence_start)}-{format_ts(evidence_end)}）。",
            )
        )
    else:
        findings.append(("No people change events detected.", "未检测到人数变化事件。"))

    return findings[:3]


def write_report_docx(report: Report, path: str, *, progress: Optional[object] = None) -> str:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise ImportError("python-docx is required for report.docx generation") from exc

    FONT_ZH = "SimSun"
    FONT_EN = "Times New Roman"
    PT_TITLE = 18
    PT_H1 = 14
    PT_BODY = 11
    PT_TABLE = 10.5
    HEADER_SHADE = "D9D9D9"

    def _progress_update(amount: float) -> None:
        if progress is not None:
            progress.update(amount)

    def set_run_font_bilingual(
        run,
        font_en: str = FONT_EN,
        font_zh: str = FONT_ZH,
        size_pt: Optional[float] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
    ) -> None:
        r = run._element
        r_pr = r.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), font_en)
        r_fonts.set(qn("w:hAnsi"), font_en)
        r_fonts.set(qn("w:eastAsia"), font_zh)
        r_fonts.set(qn("w:cs"), font_en)
        run.font.name = font_en
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic

    def _set_paragraph_spacing(paragraph, *, before_pt: float, after_pt: float, line_spacing: Optional[float]) -> None:
        p_format = paragraph.paragraph_format
        p_format.space_before = Pt(before_pt)
        p_format.space_after = Pt(after_pt)
        if line_spacing is not None:
            p_format.line_spacing = line_spacing

    def add_title(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        set_run_font_bilingual(run, size_pt=PT_TITLE, bold=True)
        _set_paragraph_spacing(paragraph, before_pt=6, after_pt=6, line_spacing=None)

    def add_h1(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        set_run_font_bilingual(run, size_pt=PT_H1, bold=True)
        _set_paragraph_spacing(paragraph, before_pt=6, after_pt=6, line_spacing=None)

    def add_para(text: str, *, italic: bool = False) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        set_run_font_bilingual(run, size_pt=PT_BODY, italic=italic)
        _set_paragraph_spacing(paragraph, before_pt=0, after_pt=3, line_spacing=1.2)

    def add_para_small(text: str, *, italic: bool = False) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        set_run_font_bilingual(run, size_pt=PT_TABLE, italic=italic)
        _set_paragraph_spacing(paragraph, before_pt=0, after_pt=3, line_spacing=1.2)

    def _shade_cell(cell, color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tc_pr.append(shd)

    def _set_cell_text(
        cell,
        text: str,
        *,
        bold: bool,
        align: str,
        size_pt: float,
        line_spacing: float,
        before_pt: float,
        after_pt: float,
    ) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        if align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(paragraph, before_pt=before_pt, after_pt=after_pt, line_spacing=line_spacing)
        run = paragraph.add_run(text)
        set_run_font_bilingual(run, size_pt=size_pt, bold=bold)

    def _disable_row_split(row) -> None:
        try:
            row.allow_break_across_pages = False
            return
        except Exception:
            pass
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    def _compute_col_widths(
        headers: List[str],
        rows: List[List[str]],
        *,
        total_cm: float,
        min_cm: float,
        max_cm_by_col: Optional[dict[int, float]],
    ) -> List[float]:
        max_cm_by_col = max_cm_by_col or {}
        col_count = len(headers)
        max_lens = [0] * col_count
        for idx, header in enumerate(headers):
            max_lens[idx] = max(max_lens[idx], len(str(header)))
        for row in rows:
            for idx, cell in enumerate(row):
                max_lens[idx] = max(max_lens[idx], len(str(cell)))
        weights = [max(1, length) for length in max_lens]
        total_weight = sum(weights) or 1
        widths = [total_cm * weight / total_weight for weight in weights]

        for _ in range(5):
            fixed = [False] * col_count
            for i, width in enumerate(widths):
                max_cm = max_cm_by_col.get(i)
                if width < min_cm:
                    widths[i] = min_cm
                    fixed[i] = True
                elif max_cm is not None and width > max_cm:
                    widths[i] = max_cm
                    fixed[i] = True
            current = sum(widths)
            diff = total_cm - current
            if abs(diff) < 1e-6:
                break
            adjustable = [i for i, is_fixed in enumerate(fixed) if not is_fixed]
            if not adjustable:
                break
            adj_weight = sum(weights[i] for i in adjustable) or 1
            for i in adjustable:
                widths[i] += diff * weights[i] / adj_weight
        return widths

    def _add_table(
        headers: List[str],
        rows: List[List[str]],
        *,
        center_cols: Optional[set[int]] = None,
        right_align_cols: Optional[set[int]] = None,
        bold_rows: Optional[set[int]] = None,
        size_pt: float = PT_TABLE,
        line_spacing: float = 1.0,
        cell_before_pt: float = 0,
        cell_after_pt: float = 0,
        table_before_pt: float = 3,
        table_after_pt: float = 6,
        total_width_cm: float = 15.9,
        min_width_cm: float = 1.2,
        max_widths: Optional[dict[int, float]] = None,
    ) -> None:
        center_cols = center_cols or set()
        right_align_cols = right_align_cols or set()
        bold_rows = bold_rows or set()
        spacing_before = document.add_paragraph()
        _set_paragraph_spacing(spacing_before, before_pt=table_before_pt, after_pt=0, line_spacing=None)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False
        for idx, header in enumerate(headers):
            _set_cell_text(
                table.rows[0].cells[idx],
                header,
                bold=True,
                align="center",
                size_pt=size_pt,
                line_spacing=line_spacing,
                before_pt=cell_before_pt,
                after_pt=cell_after_pt,
            )
            _shade_cell(table.rows[0].cells[idx], HEADER_SHADE)
        for row_idx, row in enumerate(rows):
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                if idx in right_align_cols:
                    align = "right"
                elif idx in center_cols:
                    align = "center"
                else:
                    align = "left"
                _set_cell_text(
                    row_cells[idx],
                    value,
                    bold=row_idx in bold_rows,
                    align=align,
                    size_pt=size_pt,
                    line_spacing=line_spacing,
                    before_pt=cell_before_pt,
                    after_pt=cell_after_pt,
                )
        widths = _compute_col_widths(
            headers,
            rows,
            total_cm=total_width_cm,
            min_cm=min_width_cm,
            max_cm_by_col=max_widths,
        )
        for idx, width_cm in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = Cm(width_cm)
        for row in table.rows:
            _disable_row_split(row)
        spacing_after = document.add_paragraph()
        _set_paragraph_spacing(spacing_after, before_pt=0, after_pt=table_after_pt, line_spacing=None)

    def _add_footer(video_id: str) -> None:
        section = document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"Video: {video_id} | Report {REPORT_VERSION} | Page ")
        set_run_font_bilingual(run, size_pt=PT_TABLE)

        def _add_field(instr: str) -> None:
            fld = OxmlElement("w:fldSimple")
            fld.set(qn("w:instr"), instr)
            r = OxmlElement("w:r")
            fld.append(r)
            paragraph._p.append(fld)

        _add_field("PAGE")
        run = paragraph.add_run(" of ")
        set_run_font_bilingual(run, size_pt=PT_TABLE)
        _add_field("NUMPAGES")

    document = Document()

    for section in document.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.gutter = Cm(0)

    section_no = 1

    def add_section_title(text: str) -> None:
        nonlocal section_no
        add_h1(f"{section_no}. {text}")
        section_no += 1

    add_title(f"{STRINGS['en']['report_title']} / {STRINGS['zh']['report_title']}")
    add_section_title(STRINGS["en"]["section_cover_summary"] + " / " + STRINGS["zh"]["section_cover_summary"])
    add_para(STRINGS["en"]["summary_intro"])
    add_para(STRINGS["zh"]["summary_intro"])
    add_para(
        f"People count stats: min={report.summary.min_people_count}, "
        f"max={report.summary.max_people_count}, changes={report.summary.people_change_count}"
    )
    add_para(
        f"人数统计：最小值={report.summary.min_people_count}，最大值={report.summary.max_people_count}，变化次数={report.summary.people_change_count}"
    )

    add_para(STRINGS["en"]["key_findings_title"])
    for en_line, zh_line in _build_key_findings(report):
        add_para(f"EN: {en_line}")
        add_para(f"ZH: {zh_line}")

    add_para(f"{STRINGS['en']['scope_title']} / {STRINGS['zh']['scope_title']}")
    add_para(f"Video: {report.header.source_path}")
    add_para(f"{STRINGS['en']['scope_method']}")
    add_para(f"{STRINGS['en']['scope_outputs']}")
    add_para(f"{STRINGS['en']['scope_alarm_note']}")
    add_para(f"视频来源：{report.header.source_path}")
    add_para(f"{STRINGS['zh']['scope_method']}")
    add_para(f"{STRINGS['zh']['scope_outputs']}")
    add_para(f"{STRINGS['zh']['scope_alarm_note']}")

    add_section_title(STRINGS["en"]["section_video_algo"] + " / " + STRINGS["zh"]["section_video_algo"])
    add_para(STRINGS["en"]["video_algo_intro"])
    add_para(STRINGS["zh"]["video_algo_intro"])
    add_para(STRINGS["en"]["video_algo_params_intro"])
    add_para(STRINGS["zh"]["video_algo_params_intro"])

    param_rows = [
        [
            "people_grace_s",
            f"{report.config.people_grace_s:.1f}",
            "Grace window for crew count deviation within sampling sessions.",
            "人数偏离容忍时间，采样会话内人数偏离不超过该时间不计为异常。",
        ],
        [
            "unblocked_alarm_s",
            f"{report.config.unblocked_alarm_s:.1f}",
            "Threshold for unblocked insertion accumulation to trigger an alarm.",
            "无封堵伸入累计达到该阈值触发报警。",
        ],
        [
            "gap_allow_unblocked_s",
            f"{report.config.gap_allow_unblocked_s:.1f}",
            "Gap tolerance for unblocked accumulation without reset.",
            "无封堵累计的间断容忍，短暂中断不清零。",
        ],
        [
            "sampling_min_s",
            f"{report.config.sampling_min_s:.0f}",
            "Minimum sampling duration (only when enabled).",
            "最短采样时长阈值（仅在开关启用时生效）。",
        ],
    ]
    _add_table(
        ["Parameter", "Value", "Explanation (EN)", "说明（ZH）"],
        param_rows,
        center_cols={1},
        right_align_cols={1},
        max_widths={2: 5.5, 3: 5.5},
    )
    _progress_update(15.0)

    add_section_title(STRINGS["en"]["section_sampling_overview"] + " / " + STRINGS["zh"]["section_sampling_overview"])
    add_para(STRINGS["en"]["sessions_intro"])
    add_para(STRINGS["zh"]["sessions_intro"])
    _add_table(
        ["ID", "Type", "Start", "End", "Duration(s)"],
        _session_rows(report.sessions),
        center_cols={0, 2, 3},
        right_align_cols={4},
    )
    _progress_update(15.0)

    if report.alarms:
        add_section_title(STRINGS["en"]["section_alarm_list"] + " / " + STRINGS["zh"]["section_alarm_list"])
        add_para(STRINGS["en"]["alarms_intro"])
        add_para(STRINGS["zh"]["alarms_intro"])
        alarm_rows = []
        alarm_progress_total = 30.0
        alarm_count = max(1, len(report.alarms))
        alarm_step = alarm_progress_total / alarm_count
        for alarm in report.alarms:
            duration = max(0.0, alarm.end_ts_s - alarm.start_ts_s)
            evidence_start = max(0.0, alarm.start_ts_s - 5.0)
            evidence_end = alarm.end_ts_s + 5.0
            session_id = _infer_session_id(alarm, report.sessions)
            display_type = alarm.alarm_type
            if display_type == "CREW_UNDER":
                display_type = "CREW_UNDER_2"
            elif display_type == "CREW_OVER":
                display_type = "CREW_OVER_2"
            alarm_rows.append(
                [
                    str(alarm.alarm_id),
                    display_type,
                    format_ts(alarm.start_ts_s),
                    format_ts(alarm.end_ts_s),
                    _format_duration_s(duration),
                    session_id,
                    f"{format_ts(evidence_start)} - {format_ts(evidence_end)}",
                    _explain_alarm(
                        alarm,
                        lang="en",
                        report=report,
                        session_id=session_id,
                        duration_s=duration,
                        evidence_start_s=evidence_start,
                        evidence_end_s=evidence_end,
                    ),
                    _explain_alarm(
                        alarm,
                        lang="zh",
                        report=report,
                        session_id=session_id,
                        duration_s=duration,
                        evidence_start_s=evidence_start,
                        evidence_end_s=evidence_end,
                    ),
                ]
            )
            _progress_update(alarm_step)
        _add_table(
            ["ID", "Type", "Start", "End", "Duration(s)", "Session", "Evidence", "Explanation (EN)", "说明（ZH）"],
            alarm_rows,
            center_cols={0, 2, 3, 5},
            right_align_cols={4},
            max_widths={7: 4.5, 8: 4.5, 6: 3.8},
        )
    else:
        _progress_update(30.0)

    add_section_title(STRINGS["en"]["section_people_changes"] + " / " + STRINGS["zh"]["section_people_changes"])
    add_para(STRINGS["en"]["people_changes_intro_en"])
    add_para(STRINGS["zh"]["people_changes_intro_zh"])
    _add_table(
        ["Time", "Count Change(A→B)", "Context", "Evidence(±5s)"],
        _people_change_rows(report.people_count_change_events),
        center_cols={0, 2},
        max_widths={1: 3.0, 3: 4.0},
    )

    add_section_title(STRINGS["en"]["section_people_segments"] + " / " + STRINGS["zh"]["section_people_segments"])
    add_para(STRINGS["en"]["people_segments_intro_en"])
    add_para(STRINGS["zh"]["people_segments_intro_zh"])
    _add_table(
        ["Start", "End", "Duration(s)", "People Count", "Context"],
        _people_segment_rows(report.people_count_segments),
        center_cols={0, 1, 4},
        right_align_cols={2, 3},
    )

    add_section_title(STRINGS["en"]["section_observation"])
    add_para(STRINGS["en"]["observation_intro_en"])
    add_para(STRINGS["zh"]["observation_intro_zh"])
    _add_table(
        ["ID", "Start", "End", "Duration(s)"],
        _observation_rows(report.open_no_sampling_segments),
        center_cols={0, 1, 2},
        right_align_cols={3},
        total_width_cm=15.9,
        max_widths={3: 2.5},
    )
    _progress_update(15.0)

    add_section_title(STRINGS["en"]["section_timeline"] + " / " + STRINGS["zh"]["section_timeline"])
    add_para(STRINGS["en"]["timeline_intro"])
    add_para(STRINGS["zh"]["timeline_intro"])
    timeline_rows = []
    timeline_bold_rows = set()
    timeline_progress_total = 10.0
    timeline_count = max(1, len(report.sessions) + len(report.people_count_change_events) + len(report.alarms))
    timeline_step = timeline_progress_total / timeline_count
    for session in report.sessions:
        evidence_start = max(0.0, session.start_ts_s - 5.0)
        evidence_end = session.start_ts_s + 5.0
        timeline_bold_rows.add(len(timeline_rows))
        timeline_rows.append(
            [
                format_ts(session.start_ts_s),
                "SESSION",
                f"Sampling session #{session.session_id} ({session.session_type})",
                f"{format_ts(evidence_start)} - {format_ts(evidence_end)}",
            ]
        )
        _progress_update(timeline_step)
    for change in report.people_count_change_events:
        evidence_start = max(0.0, change.change_ts_s - 5.0)
        evidence_end = change.change_ts_s + 5.0
        timeline_rows.append(
            [
                format_ts(change.change_ts_s),
                "PEOPLE CHANGE",
                f"{change.from_count}→{change.to_count}",
                f"{format_ts(evidence_start)} - {format_ts(evidence_end)}",
            ]
        )
        _progress_update(timeline_step)
    for alarm in report.alarms:
        evidence_start = max(0.0, alarm.start_ts_s - 5.0)
        evidence_end = alarm.start_ts_s + 5.0
        timeline_rows.append(
            [
                format_ts(alarm.start_ts_s),
                "ALARM",
                f"{alarm.alarm_type} #{alarm.alarm_id}",
                f"{format_ts(evidence_start)} - {format_ts(evidence_end)}",
            ]
        )
        _progress_update(timeline_step)
    if not report.sessions and not report.people_count_change_events and not report.alarms:
        _progress_update(timeline_progress_total)
    timeline_rows.sort(key=lambda row: row[0])
    _add_table(
        ["Time", "Event", "Detail", "Evidence(±5s)"],
        timeline_rows,
        center_cols={0, 1},
        bold_rows=timeline_bold_rows,
        size_pt=10.0,
        line_spacing=1.0,
        cell_before_pt=0,
        cell_after_pt=0,
        table_before_pt=3,
        table_after_pt=6,
        max_widths={3: 4.0},
    )

    add_section_title(STRINGS["en"]["section_disclaimer"] + " / " + STRINGS["zh"]["section_disclaimer"])
    add_para_small(STRINGS["en"]["disclaimer"], italic=True)
    add_para_small(STRINGS["zh"]["disclaimer"], italic=False)

    document.add_page_break()
    add_section_title(STRINGS["en"]["glossary_title"])
    glossary_rows = [
        ["Sampling session", "Sampling session (sampling operation segment)", "采样会话（采样作业段）"],
        ["Observation segment", "Observation segment", "观察区间"],
        ["People change event", "People change event", "人数变化事件"],
        ["People count segment", "People count segment", "人数阶段区间"],
        ["Evidence range", "Evidence range (+/-5s)", "证据复核范围（±5秒）"],
    ]
    _add_table(
        ["Term", "English", "中文"],
        glossary_rows,
        center_cols={0},
        max_widths={1: 6.0, 2: 6.0},
    )

    video_id = Path(report.header.source_path).name
    _add_footer(video_id)

    document.save(path)
    _progress_update(5.0)
    return path
